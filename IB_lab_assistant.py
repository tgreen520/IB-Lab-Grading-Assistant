import streamlit as st
import anthropic
import base64
import pandas as pd
import os
import zipfile
import time
import re
from docx import Document
from io import BytesIO

# --- 1. PAGE SETUP (MUST BE FIRST) ---
st.set_page_config(
    page_title="IB Lab Grader", 
    page_icon="🧪", 
    layout="wide"
)

# --- 2. CONFIGURATION & SECRETS ---
if "ANTHROPIC_API_KEY" in st.secrets:
    API_KEY = st.secrets["ANTHROPIC_API_KEY"]
elif "ANTHROPIC_API_KEY" in os.environ:
    API_KEY = os.environ.get("ANTHROPIC_API_KEY")
else:
    st.error("🚨 API Key not found!")
    st.info("On Streamlit Cloud, add your key to the 'Secrets' settings.")
    st.stop()

# --- 3. HARDCODED RUBRIC (UPDATED FOR STRICTER IB CHEMISTRY STANDARDS) ---
IB_RUBRIC = """TOTAL: 100 POINTS (10 pts per section)

1. FORMATTING (10 pts):
- Criteria: Third-person passive voice, professional tone, superscripts/subscripts used correctly.
- DEDUCTIONS: 1-2 subscript errors = -0.5 pts. 3+ errors = -1.0 pt.

2. INTRODUCTION (10 pts):
- Criteria: Clear objective, background theory, balanced equations.
- OBJECTIVE: Must be explicit. If missing, -1.0 pt.
- THEORY (STRICTER SCORING): 
  * Must explain the specific chemical theory behind the lab objective in depth.
  * **AUDIENCE STANDARD:** The background must provide sufficient detail for a chemist who is **unfamiliar with this specific experiment** to understand the context and objectives.
  * Explained fully and linked to objective: 0 deduction.
  * Discussed but generic/incomplete (not enough context for an outsider): -2.0 pts.
  * Completely missing: -3.0 pts.

3. HYPOTHESIS (10 pts):
- Criteria: Specific prediction with scientific justification.
- REQUIRED DETAILS:
  * Units for Independent Variable (IV) and Dependent Variable (DV) must be included.
  * Method of measuring the DV must be explicitly stated.
  * DV must be specific and measurable (not vague).
- DEDUCTIONS:
  * Missing units for IV/DV: -0.5 pts.
  * DV measurement is vague or measurement method is missing: -1.0 pt.

4. VARIABLES (10 pts):
- Criteria: IV, DV, 3+ Controls.
- SCORING: 
  * 10/10: All defined + explanations + units.
  - DEDUCTIONS:
  * Control variables missing (fewer than 3): -4.0
  * EXCEPTION: If the student lists specific instances (e.g. "Mass of Zinc", "Mass of Mg") as multiple IVs, do NOT deduct 4 points. Deduct ONLY 1.0 point for "Categorization Error".
  * Control variables not justified: -2.0 pts.
  * Independent variable not thoroughly explained: -1.0 pt.
  * Dependent variable not thoroughly explained: -1.0 pt.
  * Units missing for IV or DV: -0.5 pts.
  * Do not deduct points for multiple independent variables if they are derived (e.g. Temp vs 1/Temp).
  * Do not deduct for dependent variables that are derived units. 

5. PROCEDURES & MATERIALS (10 pts):
- Criteria: Numbered steps, quantities, safety, diagram.
- UNCERTAINTIES & PRECISION:
  * Uncertainties SHOULD be listed in the Materials section.
  * **SCORING:**
    - Listed in Materials: 0 Deduction.
    - Missing in Materials BUT listed in Data Section: -0.5 pts.
    - Completely missing (not in Materials OR Data): -1.0 pt.
  * Precision mismatch (uncertainty vs. measurement): -0.5 pts.
  * Diagram or photograph missing: -0.5 pt.

6. RAW DATA (10 pts):
- Criteria: Qualitative observations, tables, units, sig figs.
- REQUIREMENT: Uncertainties must be reported in Data Tables (headers or cells).
- DEDUCTIONS:
  * Missing unit uncertainties in column headers: -1.0 pt.
  * Partially included or partially correct unit uncertainties in headers: -0.5 pts.
  * Inconsistent use of significant figures: -0.5 pts.

7. DATA ANALYSIS (10 pts):
- UNCERTAINTY PROPAGATION (IB CHEMISTRY STANDARD):
  * MUST use Absolute Uncertainty (for + / -) and Percentage Uncertainty (for * / /).
  * **SCORING:**
    - **Attempted (Partial):** Any calculation OR propagation attempted, even if incomplete/incorrect: -1.0 pt.
    - **Missing:** No attempt at calculation or propagation: -2.0 pts.
- GRAPHS (Bar or Scatter allowed based on context):
  * **REQUIREMENT:** The following must be present on **ALL** relevant scatter plots.
  * Graph Missing: -2.0 pts. 
  * Axis labels missing: -1.0 pt.
  * Trendline Missing on **any** relevant graph: -1.0 pt.
  * Trendline Equation Missing on **any** relevant graph: -1.0 pt.
  * R² Value Missing on **any** relevant graph: -1.0 pt.
  * **Spectrophotometry Specific:** If this is a spectrophotometry lab, missing absorbance graph: -1.0 pt.
  * MULTIPLE TRIALS RULE: If >1 trial performed, graph MUST show AVERAGES. (If missing: -2.0 pts).

8. CONCLUSION (10 pts) [STRICT DEDUCTIONS]:
- HYPOTHESIS SUPPORT:
  * Must specifically state if data supports or refutes hypothesis using quantitative/qualitative data.
  * If missing or vague: -1.0 pt.
- STATISTICAL ANALYSIS (RELATIONSHIP):
  * Must use BOTH Coefficient of Determination (R²) AND Correlation Coefficient (r) to discuss the relationship.
        * **R (Correlation):** * **Is the R value listed?** -> If NO, deduct 1.0.
            * **Is the explanation valid?** -> If the explanation is vague OR the student confuses R with R² (e.g., "The R² shows a positive correlation"), deduct 0.5.
        * **R² (Determination):** Must explain % variation/fit. (Missing entirely -> -1.0. Vague explanation -> -0.5).
- LITERATURE COMPARISON:
  * Must compare results to published literature/accepted values.
  * If missing: -1.0 pt.
- UNCERTAINTY IMPACT (CROSS-CHECK EVALUATION): 
  * **Full Discussion:** Impact on data explained in either section: 0 deduction.
  * **Partial/Missing:** See Evaluation section logic, but ensure it is discussed somewhere.
- IV/DV RELATIONSHIP: Must explain graph trend. (If poor: -1.0)
- THEORY: Connect to chemical theory.
  * Explained fully: 0 deduction.
  * Discussed but incomplete explanation: -1.0 pt.
  * Completely missing: -2.0 pts.
- QUANTITATIVE SUPPORT: 
  * Must cite collected raw data specific numbers.
  * If NO quantitative data (neither derived nor collected) is cited: -2.0 pts.
- QUALITATIVE SUPPORT: Must cite observations. (If missing: -0.5)
- OUTLIERS/OMISSIONS: Must address data outliers or omissions. (No mention: -1.0. Mentioned but vague: -0.5).

9. EVALUATION (10 pts) [STRICT QUALITY GATES]:
- REQUIREMENT: List errors + Specific Directional Impact + Specific Improvement.
- ERROR CLASSIFICATION: Must differentiate between systematic and random errors. (Not done: -0.5).
- IMPACT SCORING (STRICT):
  * Must specifically address how EACH random and systematic error affects the measured data.
  * Impact defined for all errors: 0 deduction.
  * Impact partially addressed: -1.0 pt.
  * Impact NOT addressed: -2.0 pts.
- IMPROVEMENT SCORING:
  * Specific equipment named = 2 pts.
  * Vague ("use better scale") = 1.5 pts (Deduct 0.5).
  * Generic ("be careful") = 0 pts (Deduct 2.0).

10. REFERENCES (10 pts):
- Criteria: Citations present for all external data/images.
- SCORING:
  * 0 References provided: Score 0/10 (Deduct 10 pts).
  * Only 1 relevant source provided: Deduct 5.0 pts.
  * Only 2 relevant sources provided: Deduct 3.0 pts.
  * 3+ relevant sources provided: 0 deduction (Full points if formatted correctly).
  * Minor APA formatting errors: Deduct 1.0 pt.
"""

# --- 4. SYSTEM PROMPT (UPGRADED FOR THOROUGHNESS) ---
SYSTEM_PROMPT = """You are an expert IB Chemistry Lab Grader and Educator. 
Your goal is not just to grade, but to **teach** the student how to improve by referencing specific IB criteria.

### 🧪 SCIENTIFIC FORMATTING RULES (STRICT):
1.  **NO HTML or LATEX TAGS:** Do not use `<sup>`, `<sub>`, `$`, or markdown code blocks for chemistry.

### 🧪 YOUR OUTPUT FORMATTING (SCIENTIFIC):
1.  **USE UNICODE CHARACTERS:** Even if the student\'s text looks flat (\'cm3\'), YOUR feedback must use proper symbols.
    * *Bad:* cm^3, dm^-3, CO_2
    * *Good:* cm³, dm⁻³, CO₂, 10⁵, ±0.05
    * *Common Symbols:* ⁰ ¹ ² ³ ⁴ ⁵ ⁶ ⁷ ⁸ ⁹ ⁺ ⁻ ₀ ₁ ₂ ₃ ₄ ₅ ₆ ₇ ₈ ₉
2.  **NO FALSE ACCUSATIONS:** Do **NOT** accuse the student of using HTML tags (like `<sup>`) or LaTeX. The student is writing in Word; they are not coding. If the text looks weird, blame the PDF extractor, not the student.

### 🛡️ VOCABULARY & PHRASING IMMUNITY (DO NOT DEDUCT):
1.  **"HCl Acid" / "H₂SO₄ Acid":** While technically redundant (since \"acid\" is implied), this is a common student shorthand.
    * **Action:** Do NOT deduct points. Do NOT mention it as an error. Treat it as correct phrasing.
2.  **"Molar Mass of X":** If they say "Molecular Weight" instead of "Molar Mass," accept it.
3.  **"Experiment" vs "Investigation":** Use these interchangeably.

### 🧠 DEEP DIVE FEEDBACK PROTOCOL (MANDATORY):

1.  **THE "RULE + EVIDENCE" STANDARD:**
    * **Never** make a claim without citing the Rubric Section.
    * *Bad:* "You need more controls."
    * *Good:* "❌ **Rubric Section 4 (Variables)** requires at least 3 controlled variables to ensure a fair test. You only listed 1."

2.  **STRENGTHS = "QUOTE + CRITERIA + EFFECT":**
    * Do not just praise. You must explain **why** it was effective scientifically.
    * **Structure:** [Quote the student] -> [Cite the specific Rubric Criteria met] -> [Explain the scientific value].
    * *Example:* "✅ You successfully stated \"The uncertainty of the burette is ±0.05mL.\" This meets the **Section 5 (Precision)** requirement. Listing this allows us to propagate error correctly in the analysis."

3.  **IMPROVEMENTS = "ERROR + RULE + FIX + EDUCATIONAL REASON":**
    * **The Error:** Quote exactly what they wrote (or state \'Completely Missing\').
    * **The Rule:** Start with **"The rubric requires..."** (Do **NOT** mention specific Section numbers like \'Section 2\').
    * **The Fix:** Provide a concrete, corrected example using UNICODE.
    * **The Educational Reason:** Explain *why* this rule exists in Chemistry.
    * *Example:* "⚠️ **Error:** You wrote units as \"mol/dm3\". **Rule:** The rubric requires inverse notation. **Fix:** Write \"mol dm⁻³\". **Why?** This is the standard IUPAC notation."

### ⚖️ CALIBRATION & TIE-BREAKER STANDARDS:

1.  **THE "BENEFIT OF DOUBT" RULE:**
    * If a student\'s phrasing is clumsy but technically accurate -> **NO DEDUCTION.**
    * If a student uses the wrong vocabulary word but the concept is correct -> **-0.5 (Vague).**
    * If the text is contradictory (says X, then says Not X) -> **-1.0 (Unclear).**

2.  **THE "STRICT BINARY" DECISION TREE:**
    * **Is the Hypothesis Justification missing?** * YES -> -2.0.
        * NO, but it relies on non-scientific reasoning (e.g., "I feel like...") -> -1.0.
    * **Is the R² value on the graph?**
        * YES (Explicitly written) -> 0 deduction.
        * NO (Not visible) -> -1.0 deduction. (Do not assume it is \"implied\").

**CRITICAL INSTRUCTION:** 1. Perform ALL math calculations for ALL sections inside a single `<math_scratchpad>` block at the VERY START of your response. 
2. The user will NOT see this block (it is filtered out).
3. Do NOT include any math or deduction logic in the "OUTPUT FORMAT" sections. Only the final feedback text.

1.  INTRODUCTION (Section 2) - DEDUCTION PROTOCOL:
    * **Start at 10.0 Points.**
    * **Objective:** If Missing -> -1.0. If Vague/Implicit -> -0.5.
    * **Chemical Equation:** If Missing -> -1.0.
    * **Background Theory (STRICT QUALITY CONTROL):** * **Audience Check:** Does the background provide enough context for a **chemist unfamiliar with this specific experiment**?
        * **Missing/Irrelevant:** If the theory is missing entirely or purely historical without chemical relevance -> **-3.0 pts.**
        * **Weak/Superficial:** If the theory is present but assumes prior knowledge of the specific experiment, or is just a definition list -> **-2.0 pts.**
    * **RESTRICTIONS (Do NOT Deduct):** No deductions for citation context or inconsistent units.

2.  **CONCLUSION (Section 8) - STRICT MATH PROTOCOL:**
    * **Start at 10.0 Points.**
    * **Hypothesis Support:** Not stated? -> -1.0.
    * **Outliers/Omissions:** No mention? -> -1.0. Vague? -> -0.5.
    * **Literature Comparison:** Vague comparison (no specific values)? -> -0.5.
    * **IV/DV Trend:** Missing logic? -> -1.0.
    * **Quantitative Data:** No numbers quoted? -> -2.0.
    * **Theory:** No connection? -> -1.0.
    * **Statistics (R vs R² CHECK):**
        * **R (Correlation):** Must explain Strength AND Direction. (Missing/No explanation -> -1.0. Vague explanation -> -0.5).
        * **R² (Determination):** Must explain % variation/fit. (Missing entirely -> -2.0. Vague explanation -> -1.0).
        * **Differentiation:** Ensure student treats R and R² as separate concepts. If they mix them up, apply the "Vague" deduction for both.
    * **Focus:** Repetitive/Unfocused? -> -0.5 (Max).
    * **RESTRICTIONS (Do NOT Deduct):** NO deductions for Citations, "Internal Inconsistency", or "Data Reliability".

3.  **HYPOTHESIS (Section 3):**
    * **Justification Check:** Missing? -> -2.0. Incomplete/Vague? -> -1.0.
    * **Units Check:** Missing -> -1.0. Incomplete -> -0.5.
    * **Measurement Check:** Missing -> -1.0. Vague -> -0.5.

4.  VARIABLES (Section 4) - JUSTIFICATION PROTOCOL:
    * **Accuracy Check (NEW):** * **Swapped/Wrong Variables:** Did they list the IV as the DV (or vice versa)? Or did they list a constant as a variable? -> **-1.0 point.**
    * **Control Justification:** * No justification given for why controls were chosen? -> -2.0.
        * Partial/Vague justification? -> -1.0.
    * **DV Measurement:** Method for measuring DV is vague? -> -0.5.
    * **Identification (Missing Items):** Control variables missing? -> -1.0 per missing item. IV missing? -> -2.0. DV missing? -> -2.0.

5.  **DATA ANALYSIS (Section 7):**
    * **ALL GRAPHS CHECK:** You must check **every** scatter plot included in the report.
    * **Trendline Equation:** Not shown on **ANY** relevant graph? -> -1.0.
    * **R² Value:** Not shown on **ANY** relevant graph? -> -1.0.
    * **Trendline:** Not shown on **ANY** relevant graph? -> -1.0.
    * **Calculations:** Example calculations unclear? -> -1.0.
    * **Steps:** Calculation steps not clearly explained OR labeled? -> -0.5.

6.  **PROCEDURES (Section 5):**
    * **Diagram Check:** Diagram or photograph of experimental setup missing? -> -0.5.

7.  **EVALUATION (STRICT IMPACT AUDIT):** - **ERROR CLASSIFICATION (KEYWORD SEARCH):** Scan the text for the words "Systematic" or "Random". 
     * **If present:** Assume the student has differentiated correctly. DO NOT DEDUCT.
     * **If absent:** Deduct 0.5.
   - **MANDATORY IMPACT CHECK:** List every error the student mentions. For EACH error, verify if they explain 
     the DIRECTIONAL impact on the final calculated value (e.g., 'caused molar mass to be too high', 
     'made concentration lower than actual'). 
   - **SCORING:** If 0 errors have directional impact -> -2.0 pts. If some but not all -> -1.0 pt. 
     If all errors have direction -> No deduction.
   - In your feedback, you MUST write: 'You listed [X] errors. [Y] had explicit directional impact.' 
   - Penalize vague improvements (-0.5) or generic improvements like 'be more careful' (-2.0)."
    * **CRITICAL IMPACT AUDIT (THE "DIRECTION" CHECK - STRICTLY ENFORCE):**
        * **Step 1:** Count the TOTAL number of errors the student lists (e.g., "spilling water", "heat loss", "scale precision").
        * **Step 2:** For EACH error, search for EXPLICIT directional language about the calculated result:
            - ACCEPTABLE phrases: "made the result too high", "caused an overestimation", "led to a lower value", "increased the calculated mass", "decreased the final answer"
            - NOT ACCEPTABLE: "affected accuracy", "caused error", "impacted results", "reduced precision" (these are vague - no direction specified)
        * **Step 3:** Count how many errors have explicit directional impact.
        * **Step 4:** Apply Scoring (NO EXCEPTIONS):
            - If **ZERO** errors have directional impact explained -> **DEDUCT 2.0 points** (Max score 8.0)
            - If **SOME BUT NOT ALL** errors have directional impact -> **DEDUCT 1.0 point** (Max score 9.0)
            - If **ALL** errors have specific directional impact -> **NO DEDUCTION** (Score 10.0 possible)
        
        * **EXAMPLE GRADING:**
            - Student lists 3 errors but only explains direction for 2 of them -> DEDUCT 1.0 pt
            - Student lists 4 errors but explains direction for 0 of them -> DEDUCT 2.0 pts
            - Student lists 2 errors and explains direction for both -> NO DEDUCTION (assuming other criteria met)
    
    * **IMPROVEMENTS:** Specific equipment named? -> No deduction. Vague? -> -0.5. Generic? -> -2.0.
    
    * **MANDATORY FEEDBACK FORMAT:** In your response, you MUST explicitly state:
        - "You listed [X] total errors."
        - "Of these, [Y] had explicit directional impact on the calculated value."
        - If Y < X: "This results in a deduction of [1.0 or 2.0] points."

8.  REFERENCES (Section 10) - QUANTITY CHECK:
    * **LOGIC GATE (MANDATORY):** * **Step 1:** Search the document for a header labeled "References", "Bibliography", "Works Cited", "Sources", or "Acknowledgements".
        * **Step 2:** Is the header present?
            * **NO:** Score = 0 points.
            * **YES:** Score = **MINIMUM 4.0 POINTS.** (You are FORBIDDEN from giving 0, 1, 2, or 3 points if the section exists).
    
    * **SCORING LADDER (Only applies if header exists):**
        * **3+ Credible Sources:** 10.0 pts.
        * **2 Credible Sources:** 7.0 pts.
        * **1 Credible Source:** 5.0 pts.
        * **0 Credible Sources (e.g., all are Wikipedia, Google, or broken links):** 0.0 pts (The "Attempted" Score).
    
    * **Formatting:** Do NOT deduct for minor APA formatting errors. Deduct 0.5 points for major APA formatting errors. 
### 🧠 SCORING ALGORITHMS (STRICT ENFORCEMENT):


### 📝 FEEDBACK STYLE INSTRUCTIONS:
1. **FORMATTING:** Use <sub> and <sup> tags for chemical formulas and exponents (e.g., write H<sub>2</sub>O, 10<sup>5</sup>).
2. **AVOID ROBOTIC CHECKLISTS:** Do not use "[Yes/No]".
3. **EXPLAIN WHY:** Write 2-3 sentences for each section.
4. **TOP 3 ACTIONABLE STEPS:** You MUST provide exactly THREE specific, actionable steps at the end. These should be concrete recommendations the student can implement in their next lab report.

### OUTPUT FORMAT:
Please strictly use the following format. Do not use horizontal rules (---) between sections. Do NOT print the calculation steps here.

# 📝 SCORE: [Total Points]/100
STUDENT: [Filename]

**📊 OVERALL SUMMARY & VISUAL ANALYSIS:**
* [1-2 sentences on quality]
* [Critique of graphs/images]

**📝 DETAILED RUBRIC BREAKDOWN:**

**1. FORMATTING: [Score]/10**
* **✅ Strengths:** [Detailed explanation of tone/voice quality]
* **⚠️ Improvements:** [**MANDATORY:** "Found [X] subscript errors." (If X=1 or 2, Score **MUST** be 9.5. If X>=3, Score is 9.0 or lower).]

**2. INTRODUCTION: [Score]/10**
* **✅ Strengths:** [Detailed explanation of objective/theory coverage]
* **⚠️ Improvements:** [**CRITICAL CHECKS:** * "Objective explicit?" (-1.0 if No, -0.5 if Vague). * "Chemical Equation present?" (-1.0 if No). * "Background sufficient for unfamiliar audience?" (-2.0 if generic, -3.0 if missing/irrelevant). Ensure theory explains the chemistry clearly to an outsider.]

**3. HYPOTHESIS: [Score]/10**
* **✅ Strengths:** [Quote prediction and praise the scientific reasoning]
* **⚠️ Improvements:** [**CRITICAL CHECKS:**
* "Justification: [Present/Missing/Vague]" (-2.0 if missing, -1.0 if vague/incomplete).
* "Units for IV/DV: [Present/Missing]" (-1.0 if missing, -0.5 if partial).
* "DV Measurement Description: [Specific/Vague/Missing]" (-1.0 if missing, -0.5 if vague).]

**4. VARIABLES: [Score]/10**
* **✅ Strengths:** [**LIST:** "Identified IV: [X], DV: [Y], Controls: [A, B, C]" and comment on clarity.]
* **⚠️ Improvements:** [If DV measurement is vague, state: "The method for measuring the DV was vague (-0.5 pts)." Suggest specific improvement.]

**5. PROCEDURES: [Score]/10**
* **✅ Strengths:** [Comment on reproducibility and safety details]
* **⚠️ Improvements:** [**DIAGRAM CHECK:** "Diagram of experimental setup included?" (-0.5 if missing). Identify exactly which step is vague and how to fix it.]

**6. RAW DATA: [Score]/10**
* **✅ Strengths:** [Comment on data organization and unit clarity]
* **⚠️ Improvements:** [Quote values with wrong units/sig figs and explain the correct format. Comment on inconsistent sig fig reporting for measuring tools.]

**7. DATA ANALYSIS: [Score]/10**
* **✅ Strengths:** [Summarize the calculation process. If Graph is perfect, mention that the scatterplot, equation, and labels are all correct here.]
* **⚠️ Improvements:** [**GRAPH AUDIT:** "Trendline: [Present/Missing]" (-1.0). "Equation: [Present/Missing]" (-1.0). "R² Value: [Present/Missing]" (-1.0). Check **ALL** relevant scatter plots. If even one is missing these elements, deduct.
**CALCULATION AUDIT:** "Example calculations were [Clear/Unclear]." (If unclear, -1.0 pts). "Calculation steps were [Clearly Explained/Not Labeled or Explained]." (If not labeled/explained, -0.5 pts).]

**8. CONCLUSION (10 pts) [STRICT DEDUCTIONS]:
- HYPOTHESIS SUPPORT: Must indicate if data supports hypothesis. (If missing: -1.0).
- OUTLIERS/OMISSIONS: Must address data outliers or omissions. (No mention: -1.0. Mentioned but vague: -0.5).
- IV/DV RELATIONSHIP: Must explain graph trend. (If poor: -1.0).
- THEORY: Connect to chemical theory. (If missing: -1.0).
- QUANTITATIVE SUPPORT: Must cite specific numbers. (If missing: -2.0).
- QUALITATIVE SUPPORT: Must cite observations. (If missing: -0.5).
- LITERATURE COMPARISON: If comparison to literature is vague (no specific values), -0.5 pt.
- STATISTICS (CORRELATION COEFFICIENT - R):
  * Requirement: Must explicitly list the R value. (Missing: -1.0).
  * Explanation: Must explain Strength & Direction.
  * DEDUCTION: If R is present but explanation is vague OR student confuses R with R² (e.g., uses R² to describe direction) = -0.5 pts.
- STATISTICS (R² - DETERMINATION):
  * Requirement: Must explain Fit/Variability. (Missing: -1.0. Vague: -0.5).
- NOTE: Do NOT deduct for "Internal Inconsistency" or Citations here.

**9. EVALUATION: [Score]/10**
* **✅ Strengths:** [**LIST:** "You identified: [Error 1], [Error 2]..." and comment on depth.]
* **⚠️ Improvements:** [**ERROR CLASSIFICATION:** "You did not differentiate between systematic and random errors. (-0.5 pt)" OR "You successfully distinguished systematic from random errors."
**IMPACT/IMPROVEMENT AUDIT:** * "You listed [X] errors but only provided specific directional impacts for [Y] of them. (-1 pt)"
  * "Improvements were listed but were slightly vague (e.g., did not name specific equipment). (-0.5 pt)" ]

**10. REFERENCES: [Score]/10**
* **✅ Strengths:** [**MANDATORY:** "Counted [X] credible sources."]
* **⚠️ Improvements:** [**QUANTITY CHECK:** "Only found [X] sources." (If 1 source -> Score 5.0. If 2 sources -> Score 7.0. If References are attempted -> Score 4.0). **FORMATTING:** "APA Formatting Check: [Correct/Incorrect]" (-0.5 if incorrect).]

**💡 TOP 3 ACTIONABLE STEPS FOR NEXT TIME:**
1. [Step 1 - Specific and concrete recommendation]
2. [Step 2 - Specific and concrete recommendation]
3. [Step 3 - Specific and concrete recommendation]
"""

# Initialize Session State
if 'saved_sessions' not in st.session_state:
    st.session_state.saved_sessions = {}
if 'current_results' not in st.session_state:
    st.session_state.current_results = []
if 'current_session_name' not in st.session_state:
    st.session_state.current_session_name = "New Grading Session"
if 'autosave_dir' not in st.session_state:
    st.session_state.autosave_dir = "autosave_feedback_ib"

client = anthropic.Anthropic(api_key=API_KEY)

# --- 5. HELPER FUNCTIONS ---
def encode_file(uploaded_file):
    try:
        uploaded_file.seek(0)
        return base64.b64encode(uploaded_file.read()).decode('utf-8')
    except Exception as e:
        st.error(f"Error encoding file: {e}")
        return None

def get_media_type(filename):
    ext = filename.lower().split('.')[-1]
    media_types = {
        'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
        'gif': 'image/gif', 'webp': 'image/webp', 'pdf': 'application/pdf'
    }
    return media_types.get(ext, 'image/jpeg')

def get_para_text_with_formatting(para):
    """Iterate through runs to capture subscript/superscript formatting."""
    text_parts = []
    for run in para.runs:
        text = run.text
        if run.font.subscript:
            text = f"<sub>{text}</sub>"
        elif run.font.superscript:
            text = f"<sup>{text}</sup>"
        text_parts.append(text)
    return "".join(text_parts)

def extract_text_from_docx(file):
    try:
        file.seek(0) 
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(get_para_text_with_formatting(para))
        if doc.tables:
            full_text.append("\n--- DETECTED TABLES ---\n")
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_content = []
                        for para in cell.paragraphs:
                            cell_content.append(get_para_text_with_formatting(para))
                        row_text.append(" ".join(cell_content).strip())
                    full_text.append(" | ".join(row_text))
                full_text.append("\n") 
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading .docx file: {e}"

def extract_images_from_docx(file):
    images = []
    try:
        file.seek(0)
        with zipfile.ZipFile(file) as z:
            for filename in z.namelist():
                if filename.startswith('word/media/') and filename.split('.')[-1].lower() in ['png', 'jpg', 'jpeg', 'gif']:
                    img_data = z.read(filename)
                    b64_img = base64.b64encode(img_data).decode('utf-8')
                    ext = filename.split('.')[-1].lower()
                    images.append({
                        "type": "image",
                        "source": {
                            "type": "base64", 
                            "media_type": f"image/{'jpeg' if ext=='jpg' else ext}", 
                            "data": b64_img
                        }
                    })
    except Exception as e:
        print(f"Image extraction failed: {e}")
    return images

def process_uploaded_files(uploaded_files):
    final_files = []
    IGNORED_FILES = {'.ds_store', 'desktop.ini', 'thumbs.db', '__macosx'}
    VALID_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'gif', 'webp', 'docx'}
    file_counts = {"pdf": 0, "docx": 0, "image": 0, "ignored": 0}

    for file in uploaded_files:
        file_name_lower = file.name.lower()
        if file_name_lower in IGNORED_FILES or file_name_lower.startswith('._'):
            continue
        if file_name_lower.endswith('.zip'):
            try:
                with zipfile.ZipFile(file) as z:
                    for filename in z.namelist():
                        clean_name = filename.lower()
                        if any(x in clean_name for x in IGNORED_FILES) or filename.startswith('.'): continue
                        ext = clean_name.split('.')[-1]
                        if ext in VALID_EXTENSIONS:
                            file_bytes = z.read(filename)
                            virtual_file = BytesIO(file_bytes)
                            virtual_file.name = os.path.basename(filename)
                            final_files.append(virtual_file)
                            if ext == 'docx': file_counts['docx'] += 1
                            elif ext == 'pdf': file_counts['pdf'] += 1
                            else: file_counts['image'] += 1
            except Exception as e:
                st.error(f"Error unzipping {file.name}: {e}")
        else:
            ext = file_name_lower.split('.')[-1]
            if ext in VALID_EXTENSIONS:
                final_files.append(file)
                if ext == 'docx': file_counts['docx'] += 1
                elif ext == 'pdf': file_counts['pdf'] += 1
                else: file_counts['image'] += 1
            else:
                file_counts['ignored'] += 1
    return final_files, file_counts

def clean_hidden_math(text):
    """Removes the <math_scratchpad> and <<<MATH: ... >>> blocks from the AI output."""
    # Remove XML style scratchpad
    text = re.sub(r'<math_scratchpad>.*?</math_scratchpad>', '', text, flags=re.DOTALL)
    # Remove old style blocks just in case
    text = re.sub(r'<<<MATH:.*?>>>', '', text, flags=re.DOTALL)
    return text.strip()

def recalculate_total_score(text):
    """
    Robustly parses section scores even if the AI uses inconsistent bolding/markdown,
    sums them up, and overwrites the header score.
    """
    try:
        # 1. ROBUST PATTERN: 
        # Matches: "1. SECTION NAME: [**]9.5[**]/10"
        pattern = r"\d+\..+?:[^0-9\n]*?(\d+\.?\d*)[^0-9\n]*?/\s*10"
        
        matches = re.findall(pattern, text)
        
        if matches:
            # Convert extracted strings to floats
            scores = []
            for s in matches:
                try:
                    scores.append(float(s))
                except ValueError:
                    continue
            
            # 2. SANITY CHECK: expected 10 sections
            if len(scores) != 10:
                print(f"DEBUG: Warning - Found {len(scores)} section scores (expected 10).")
            
            total_score = sum(scores)
            
            # Format: Integer if whole number, else 1 decimal place
            if total_score.is_integer():
                total_score_str = f"{int(total_score)}"
            else:
                total_score_str = f"{total_score:.1f}"
            
            print(f"DEBUG: Recalculated Total: {total_score_str} (from {scores})")
            
            # 3. ROBUST HEADER REPLACEMENT
            # Look for "SCORE:" followed by any junk, then the old score, then "/100"
            header_pattern = r"(#\s*[🔍📝]?\s*SCORE\s*:\s*)([\d\.]+)(\s*/\s*100)"
            
            if re.search(header_pattern, text, re.IGNORECASE):
                text = re.sub(
                    header_pattern, 
                    f"\\g<1>{total_score_str}\\g<3>", 
                    text, 
                    count=1,
                    flags=re.IGNORECASE
                )
            else:
                # If header is missing or formatted oddly, force prepend it
                text = f"# 📝 SCORE: {total_score_str}/100\n\n" + text
                
        else:
            print("DEBUG: No section scores found to recalculate.")
            
    except Exception as e:
        print(f"ERROR in recalculate_total_score: {e}")
    
    return text

def parse_feedback_for_csv(text):
    data = {}
    clean_text = re.sub(r'[*#]', '', text) 
    try:
        summary_match = re.search(r"OVERALL SUMMARY.*?:\s*\n(.*?)(?=1\.|DETAILED)", clean_text, re.DOTALL | re.IGNORECASE)
        if summary_match:
            raw_summary = summary_match.group(1).strip()
            data["Overall Summary"] = re.sub(r'[\r\n]+', ' ', raw_summary)
        else:
            data["Overall Summary"] = "Summary not found"
    except Exception as e:
        data["Overall Summary"] = f"Parsing Error: {e}"

    sections = re.findall(r"(\d+)\.\s+([A-Za-z\s&]+):\s+([\d\.]+)/10\s*\n(.*?)(?=\n\d+\.|\Z|💡)", clean_text, re.DOTALL)
    for _, name, score, content in sections:
        col_name = name.strip().title()
        data[f"{col_name} Score"] = score
        cleaned_feedback = re.sub(r'[\r\n]+', ' ', content.strip())
        data[f"{col_name} Feedback"] = cleaned_feedback
    return data

def audit_score_with_ai(client, feedback_text):
    """
    Uses a second AI call to strictly parse and sum the section scores.
    """
    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514", # Or your preferred model
            max_tokens=50,
            temperature=0,
            system="You are a math auditor. Your ONLY job is to extract section scores from text and sum them.",
            messages=[
                {"role": "user", "content": f"""
                Read the following grading feedback. 
                1. Identify every section score (out of 10).
                2. Sum them up to get the true total.
                3. Return ONLY the final integer number (e.g., 85). Do not write any other text.

                FEEDBACK TO AUDIT:
                {feedback_text}
                """}
            ]
        )
        # Extract the number from the response
        true_total = response.content[0].text.strip()
        # Ensure we only got digits
        if true_total.isdigit():
            return int(true_total)
        else:
            return None
    except Exception as e:
        print(f"Math Audit Error: {e}")
        return None

# --- NEW FUNCTION: AUTOSAVE INDIVIDUAL REPORT ---
def autosave_report(item, autosave_dir):
    """Save individual report as Word doc and append to CSV immediately after grading."""
    try:
        # --- FIX: FORCE FOLDER CREATION ---
        if not os.path.exists(autosave_dir):
            os.makedirs(autosave_dir)
        # ----------------------------------
        # 1. Save Word Document
        doc = Document()
        write_markdown_to_docx(doc, item['Feedback'])
        safe_filename = os.path.splitext(item['Filename'])[0] + "_Feedback.docx"
        doc_path = os.path.join(autosave_dir, safe_filename)
        doc.save(doc_path)
        
        # 2. Append to CSV (or create if doesn't exist)
        csv_path = os.path.join(autosave_dir, "gradebook.csv")
        
        # Parse feedback into row data
        row_data = {
            "Filename": item['Filename'],
            "Overall Score": item['Score']
        }
        feedback_data = parse_feedback_for_csv(item['Feedback'])
        row_data.update(feedback_data)
        
        # Check if CSV exists
        if os.path.exists(csv_path):
            existing_df = pd.read_csv(csv_path)
            # Remove duplicate if re-grading same file
            existing_df = existing_df[existing_df['Filename'] != item['Filename']]
            new_df = pd.concat([existing_df, pd.DataFrame([row_data])], ignore_index=True)
        else:
            new_df = pd.DataFrame([row_data])
        
        new_df.to_csv(csv_path, index=False, encoding='utf-8-sig')
        
        return True
    except Exception as e:
        print(f"Autosave failed for {item['Filename']}: {e}")
        return False

def grade_submission(file, model_id):
    ext = file.name.split('.')[-1].lower()
    
    # Updated Prompt Construction for Educational Depth
    user_instructions = (
        "Please grade this lab report based on the provided rubric.\n"
        "🚨 **INSTRUCTION FOR FEEDBACK DEPTH:**\n"
        "1. **BE SPECIFIC:** Do not be vague. If you deduct points, you must explain exactly **WHY**.\n"
        "2. **BE EDUCATIONAL:** Explain the scientific reason behind the rules.\n"
        "3. **PHRASING:** When citing rules, simply say **'The rubric requires...'**. Do NOT cite specific section numbers (e.g. do NOT say 'Section 4 requires...').\n"
        "\n⚠️ **CRITICAL RUBRIC UPDATES TO ENFORCE:**\n"
        "4. **FORMATTING:** \n"
        "   - **Redundancy:** Do NOT deduct for terms like 'HCl acid' or 'Na salt'. Ignore this redundancy completely.\n"
        "1. **MATERIALS:** Look for uncertainty values (±). If missing in Materials but present in Data -> -0.5 only.\n"
        "2. **DATA ANALYSIS:** \n"
        "   - **Attempt Rule:** If they attempted ANY uncertainty math (even if wrong) -> Deduct 1.0. Only deduct 2.0 if completely missing.\n"
        "   - **Derived Independent Variables:** If they list 'Temp' and '1/Temp' as two IVs, this is CORRECT. Do not deduct.\n"
        "3. **VARIABLES:** \n"
        "   - **Categorization Error:** If they list specific instances (e.g. Zinc, Mg) instead of a category (Type of Metal) -> Deduct 1.0 (Categorization), NOT 4.0 (Missing Controls).\n"
        "   - **Derived Independent Variables:** Do NOT deduct points if the student lists multiple Independent Variables where the extra ones are mathematically derived from the main IV (e.g., 'Temperature' and '1/Temperature', or 'Concentration' and 'Natural Log of Concentration'). Treat this as a single, valid IV setup.\n"
        "   - **Derived Dependent Variables:** Do NOT deduct points if the student lists multiple Dependent Variables where the extra ones are mathematically derived from the main IV (e.g., 'Temperature' and '1/Temperature', or 'Concentration' and 'Natural Log of Concentration'). Treat this as a single, valid IV setup.\n"
        "   - **Misidentified IVs (Categorization Error):** If a student lists specific instances (e.g., 'Mass of Magnesium', 'Mass of Zinc') as multiple Independent Variables instead of the general category (e.g., 'Type of Metal'), deduct **ONLY 1.0 point** for 'Improper Variable Classification'. Do **NOT** deduct 4.0 points. Do **NOT** treat this as 'Missing Control Variables'.\n"
    )

    if ext == 'docx':
        text_content = extract_text_from_docx(file)
        if len(text_content.strip()) < 50:
            text_content += "\n\n[SYSTEM NOTE: Very little text extracted.]"
            
        prompt_text = (
            f"{user_instructions}\n"
            "--- RUBRIC START ---\n" + IB_RUBRIC + "\n--- RUBRIC END ---\n\n"
            "STUDENT TEXT:\n" + text_content
        )
        
        user_message = [{"type": "text", "text": prompt_text}]
        images = extract_images_from_docx(file)
        if images:
            user_message.extend(images)
    else:
        base64_data = encode_file(file)
        if not base64_data: return "Error processing file."
        media_type = get_media_type(file.name)
        
        prompt_text = (
            f"{user_instructions}\n"
            "--- RUBRIC START ---\n" + IB_RUBRIC + "\n--- RUBRIC END ---\n"
        )
        
        user_message = [
            {"type": "text", "text": prompt_text},
            {"type": "document" if media_type == 'application/pdf' else "image",
             "source": {"type": "base64", "media_type": media_type, "data": base64_data}}
        ]

    max_retries = 5 
    retry_delay = 5 
    
    for attempt in range(max_retries):
        try:
            response = client.messages.create(
                model=model_id,
                max_tokens=3500,
                temperature=0.0,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_message}]
            )
            raw_text = response.content[0].text
            
            # 1. Clean the Hidden Math (so user doesn't see it)
            clean_text = clean_hidden_math(raw_text)
            
            # 2. Recalculate Total (Just in case)
            final_text = recalculate_total_score(clean_text)
            
            return final_text
            
        except (anthropic.RateLimitError, anthropic.APIStatusError) as e:
            if isinstance(e, anthropic.APIStatusError) and e.status_code == 529:
                time.sleep(retry_delay * (attempt + 1))
                continue
            if isinstance(e, anthropic.RateLimitError):
                time.sleep(retry_delay * (attempt + 1))
                continue
            return f"⚠️ Error: {str(e)}"
        except Exception as e:
            return f"⚠️ Error: {str(e)}"

def parse_score(text):
    try:
        # Robust Match: Handles 📝, 🔍, or no emoji at all
        match = re.search(r"(?:#\s*[🔍📝]?\s*)?SCORE:\s*([\d\.]+)/100", text, re.IGNORECASE)
        if match: 
            return match.group(1).strip()
    except Exception as e:
        print(f"Error parsing score: {e}")
    return "N/A"

# --- WORD FORMATTER (Strict Symbol Cleaning) ---
def write_markdown_to_docx(doc, text):
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue # SKIP EMPTY LINES FOR CONTINUOUS FLOW
        
        # 1. Handle Score Header & Student Name (Larger - Level 2)
        if line.startswith('# ') or line.startswith('STUDENT:'): 
            clean = line.replace('# ', '').replace('*', '').strip()
            # Changed from Level 4 (Small) to Level 2 (Large)
            doc.add_heading(clean, level=2) 
            continue
        
        # 2. Handle H3 (### ) - CLEANED
        if line.startswith('### '):
            clean = line.replace('### ', '').replace('*', '').strip()
            doc.add_heading(clean, level=3)
            continue
        
        # 3. Handle H2 (## ) - CLEANED
        if line.startswith('## '): 
            clean = line.replace('## ', '').replace('*', '').strip()
            doc.add_heading(clean, level=2)
            continue
        
        # 4. REMOVE SEPARATORS
        if line.startswith('---') or line.startswith('___'):
            continue

        # 5. Handle Bullets (* or -) - CLEANED
        if line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            content = line[2:] 
        else:
            p = doc.add_paragraph()
            content = line

        # 6. Handle Bold (**text**) - CLEANED
        parts = re.split(r'(\*\*.*?\*\*)', content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                clean_text = part[2:-2].replace('*', '') # Strip any lingering asterisks
                run = p.add_run(clean_text)
                run.bold = True
            else:
                p.add_run(part.replace('*', '')) # Strip lingering asterisks

def create_master_doc(results, session_name):
    doc = Document()
    # REMOVED SESSION HEADER
    # doc.add_heading(f"Lab Report Grades: {session_name}", 0) 
    for item in results:
        # REMOVED FILENAME HEADER (Starts with Score + Student Name)
        write_markdown_to_docx(doc, item['Feedback'])
        doc.add_page_break()
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_zip_bundle(results):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as z:
        for item in results:
            doc = Document()
            # REMOVED FEEDBACK HEADER
            write_markdown_to_docx(doc, item['Feedback'])
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            safe_name = os.path.splitext(item['Filename'])[0] + "_Feedback.docx"
            z.writestr(safe_name, doc_buffer.getvalue())
    return zip_buffer.getvalue()

# --- NEW: AUTOSAVE INDIVIDUAL REPORT ---
def autosave_report(item, autosave_dir):
    """Save individual report as Word doc and append to CSV immediately after grading."""
    try:
        # 1. Save Word Document
        doc = Document()
        write_markdown_to_docx(doc, item['Feedback'])
        safe_filename = os.path.splitext(item['Filename'])[0] + "_Feedback.docx"
        doc_path = os.path.join(autosave_dir, safe_filename)
        doc.save(doc_path)
        
        # 2. Append to CSV (or create if doesn't exist)
        csv_path = os.path.join(autosave_dir, "gradebook.csv")
        
        # Parse feedback into row data
        row_data = {
            "Filename": item['Filename'],
            "Overall Score": item['Score']
        }
        feedback_data = parse_feedback_for_csv(item['Feedback'])
        row_data.update(feedback_data)
        
        # Check if CSV exists
        if os.path.exists(csv_path):
            existing_df = pd.read_csv(csv_path)
            # Remove duplicate if re-grading same file
            existing_df = existing_df[existing_df['Filename'] != item['Filename']]
            new_df = pd.concat([existing_df, pd.DataFrame([row_data])], ignore_index=True)
        else:
            new_df = pd.DataFrame([row_data])
        
        new_df.to_csv(csv_path, index=False, encoding='utf-8-sig')
        
        return True
    except Exception as e:
        print(f"Autosave failed for {item['Filename']}: {e}")
        return False

def display_results_ui():
    if not st.session_state.current_results:
        return

    st.divider()
    st.subheader(f"📊 Results: {st.session_state.current_session_name}")
    
    # --- EXPANDED CSV LOGIC WITH SORTING ---
    results_list = []
    for item in st.session_state.current_results:
        row_data = {
            "Filename": item['Filename'],
            "Overall Score": item['Score']
        }
        feedback_data = parse_feedback_for_csv(item['Feedback'])
        row_data.update(feedback_data)
        results_list.append(row_data)
        
    csv_df = pd.DataFrame(results_list)
    
    # Sort columns to put Filename/Score/Summary first
    cols = list(csv_df.columns)
    priority = ['Filename', 'Overall Score', 'Overall Summary']
    remaining = [c for c in cols if c not in priority]
    # Simple logic to keep section score/feedback adjacent
    remaining.sort(key=lambda x: (x.split(' ')[0], 'Feedback' in x)) 
    
    final_cols = [c for c in priority if c in cols] + remaining
    csv_df = csv_df[final_cols]
    
    csv_data = csv_df.to_csv(index=False).encode('utf-8-sig') 
    
    master_doc_data = create_master_doc(st.session_state.current_results, st.session_state.current_session_name)
    zip_data = create_zip_bundle(st.session_state.current_results)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("📄 Google Docs Compatible (.docx)", master_doc_data, f'{st.session_state.current_session_name}_Docs.docx', "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        st.caption("Upload to Drive -> Open as Google Doc")
    with col2:
        st.download_button("📦 Student Bundle (.zip)", zip_data, f'{st.session_state.current_session_name}_Students.zip', "application/zip", use_container_width=True)
    with col3:
        st.download_button("📊 Detailed CSV Export", csv_data, f'{st.session_state.current_session_name}_Detailed.csv', "text/csv", use_container_width=True)
        st.caption("Includes separate columns for every section score and comment.")

    # --- NEW: AUTOSAVE FOLDER ACCESS ---
    st.divider()
    st.info("💾 **Auto-saved files:** Individual feedback documents and gradebook are being saved to the `autosave_feedback` folder as grading progresses.")
    
    autosave_path = st.session_state.autosave_dir
    if os.path.exists(autosave_path):
        csv_autosave = os.path.join(autosave_path, "gradebook.csv")
        if os.path.exists(csv_autosave):
            with open(csv_autosave, 'rb') as f:
                st.download_button(
                    "📥 Download Auto-saved Gradebook (CSV)",
                    f.read(),
                    "autosaved_gradebook.csv",
                    "text/csv",
                    use_container_width=True
                )
        
        # Create zip of all autosaved Word docs
        autosave_files = [f for f in os.listdir(autosave_path) if f.endswith('.docx')]
        if autosave_files:
            zip_autosave = BytesIO()
            with zipfile.ZipFile(zip_autosave, 'w', zipfile.ZIP_DEFLATED) as z:
                for filename in autosave_files:
                    file_path = os.path.join(autosave_path, filename)
                    z.write(file_path, filename)
            
            st.download_button(
                "📥 Download All Auto-saved Word Docs (.zip)",
                zip_autosave.getvalue(),
                "autosaved_feedback.zip",
                "application/zip",
                use_container_width=True
            )

    # 1. Show the Gradebook Table
    st.write("### 🏆 Gradebook")
    st.dataframe(csv_df, use_container_width=True)
    
    # 2. Show the Feedback (Stacked directly below, no hiding!)
    st.write("### 📝 Detailed Feedback History")
    
    # We use reversed() so the newest file is always at the top
    for item in reversed(st.session_state.current_results):
        with st.expander(f"📄 {item['Filename']} (Score: {item['Score']})"):
            st.markdown(item['Feedback'])

# --- 6. SIDEBAR ---
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # UPDATED DEFAULT MODEL ID
    user_model_id = st.text_input(
        "🤖 Model ID", 
        value="claude-sonnet-4-20250514", 
        help="Change this if you have a specific Beta model or newer ID"
    )
    
    st.divider()
    st.header("💾 History Manager")
    save_name = st.text_input("Session Name", placeholder="e.g. Period 3 - Kinetics")
    if st.button("💾 Save Session"):
        if st.session_state.current_results:
            st.session_state.saved_sessions[save_name] = st.session_state.current_results
            st.success(f"Saved '{save_name}'!")
        else:
            st.warning("No results to save yet.")
            
    if st.session_state.saved_sessions:
        st.divider()
        st.subheader("📂 Load Session")
        selected_session = st.selectbox("Select Batch", list(st.session_state.saved_sessions.keys()))
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Load"):
                st.session_state.current_results = st.session_state.saved_sessions[selected_session]
                st.session_state.current_session_name = selected_session
                st.rerun()
        with col2:
            if st.button("🗑️ Delete"):
                del st.session_state.saved_sessions[selected_session]
                st.rerun()

    st.divider() 
    
    with st.expander("View Grading Criteria"):
        # CHANGED FROM st.text(PRE_IB_RUBRIC) TO st.text(IB_RUBRIC)
        st.text(IB_RUBRIC)

# --- 7. MAIN INTERFACE ---
st.title("🧪 IB Lab Grader")
st.caption(f"Current Session: **{st.session_state.current_session_name}**")

st.info("💡 **Tip:** To upload a folder, open it, press `Ctrl+A` (Select All), and drag everything here.")

raw_files = st.file_uploader(
    "📂 Upload Reports (PDF, Word, Images, ZIP)", 
    type=['pdf', 'docx', 'png', 'jpg', 'jpeg', 'zip'], 
    accept_multiple_files=True
)

processed_files = []
if raw_files:
    processed_files, counts = process_uploaded_files(raw_files)
    if len(processed_files) > 0:
        st.success(f"✅ Found **{len(processed_files)}** valid reports.")
        st.caption(f"📄 PDFs: {counts['pdf']} | 📝 Word Docs: {counts['docx']} | 🖼️ Images: {counts['image']}")
        if counts['ignored'] > 0:
            st.warning(f"⚠️ {counts['ignored']} files were ignored (unsupported format).")
    else:
        if raw_files:
            st.warning("No valid PDF, Word, or Image files found.")

if st.button("🚀 Grade Reports", type="primary", disabled=not processed_files):
    
    st.write("---")
    progress = st.progress(0)
    status_text = st.empty()
    live_results_table = st.empty()
    
    # NEW: Placeholder for cumulative feedback display (cleared and rewritten each iteration)
    st.subheader("📋 Live Grading Feedback")
    feedback_placeholder = st.empty()
    
    # Initialize Session State list if not present
    if 'current_results' not in st.session_state:
        st.session_state.current_results = []
    
    # Create a set of already graded filenames for quick lookup
    existing_filenames = {item['Filename'] for item in st.session_state.current_results}
    
    for i, file in enumerate(processed_files):
        # 1. SMART RESUME CHECK: Skip if already graded
        if file.name in existing_filenames:
            status_text.info(f"↩ Skipping **{file.name}** (Already Graded)")
            time.sleep(0.5) # Brief pause for visual feedback
            progress.progress((i + 1) / len(processed_files))
            continue

        # 2. GRADING LOGIC
        status_text.markdown(f"**Grading:** `{file.name}` ({i+1}/{len(processed_files)})...")
        
        try:
            # Polite delay to prevent API overloading
            time.sleep(2) 
            
            feedback = grade_submission(file, user_model_id) # PASSING USER MODEL ID
            score = parse_score(feedback)
            
            # 3. IMMEDIATE SAVE TO SESSION STATE
            new_entry = {
                "Filename": file.name,
                "Score": score,
                "Feedback": feedback
            }
            
            st.session_state.current_results.append(new_entry)
            
            # 4. AUTOSAVE TO DISK (NEW - CRITICAL FOR RECOVERY)
            autosave_success = autosave_report(new_entry, st.session_state.autosave_dir)
            if autosave_success:
                status_text.success(f"✅ **{file.name}** graded & auto-saved! (Score: {score}/100)")
            else:
                status_text.warning(f"⚠️ **{file.name}** graded but autosave failed (Score: {score}/100)")
            
            # Update the existing set so duplicates within the same batch run are also caught
            existing_filenames.add(file.name)
            
            # 5. LIVE TABLE UPDATE
            df_live = pd.DataFrame(st.session_state.current_results)
            live_results_table.dataframe(df_live[["Filename", "Score"]], use_container_width=True)
            
            # 6. UPDATED: SINGLE COPY CUMULATIVE FEEDBACK DISPLAY
            # Clear and rewrite the entire feedback section to avoid duplicates
            with feedback_placeholder.container():
                for idx, item in enumerate(st.session_state.current_results):
                    # Start expanded for most recent, collapsed for older ones
                    is_most_recent = (idx == len(st.session_state.current_results) - 1)
                    with st.expander(f"📄 {item['Filename']} (Score: {item['Score']}/100)", expanded=is_most_recent):
                        st.markdown(item['Feedback'])
            
        except Exception as e:
            st.error(f"❌ Error grading {file.name}: {e}")
            
        progress.progress((i + 1) / len(processed_files))
        

    status_text.success("✅ Grading Complete! All reports auto-saved.")
    progress.empty()
    
    # Show message about autosave location
    st.info(f"💾 **Backup Location:** All feedback has been saved to `{st.session_state.autosave_dir}/` folder. You can download individual files or the full gradebook below.")

# --- 8. PERSISTENT DISPLAY ---
if st.session_state.current_results:
    display_results_ui()